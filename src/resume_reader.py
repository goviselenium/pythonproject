import os
from pathlib import Path

try:
    import docx
except ImportError:
    docx = None

def read_resume(file_path: str | Path) -> str:
    """
    Reads the content of a master resume file.
    Supports .txt and .docx files.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found at: {path}")

    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_txt(path)
    elif ext == ".docx":
        return _read_docx(path)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Only .txt and .docx are supported.")

def _read_txt(path: Path) -> str:
    """Reads a plain text file with UTF-8 encoding (falling back to CP1252 if needed)."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1252") as f:
            return f.read().strip()

def _read_docx(path: Path) -> str:
    """Reads a DOCX file using python-docx."""
    if docx is None:
        raise ImportError(
            "The 'python-docx' library is required to read .docx files. "
            "Please run 'pip install python-docx' or use a .txt resume file."
        )
    
    try:
        doc = docx.Document(path)
        content_parts = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text.strip())
                
        # Read tables to capture any tabular structured info in resume
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # De-duplicate adjacent identical cells (often created by merging)
                clean_row = []
                for cell in row_cells:
                    if not clean_row or clean_row[-1] != cell:
                        clean_row.append(cell)
                if clean_row:
                    content_parts.append(" | ".join(clean_row))
                    
        return "\n\n".join(content_parts)
    except Exception as e:
        raise RuntimeError(f"Error reading DOCX file at {path}: {str(e)}")
